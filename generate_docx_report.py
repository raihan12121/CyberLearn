import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Set shading color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Set inner cell margins (in twips: 20 twips = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def style_table(table, col_widths, is_header_dark=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        # Prevent row split across pages
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        # Header repeat
        if i == 0 and is_header_dark:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for j, cell in enumerate(row.cells):
            set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if j < len(col_widths):
                cell.width = Inches(col_widths[j])
            if i == 0 and is_header_dark:
                set_cell_background(cell, "1F4E79") # Deep Blue
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif i % 2 == 1:
                set_cell_background(cell, "F2F5F9") # Soft Blue Tint
            else:
                set_cell_background(cell, "FFFFFF")

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    cell.width = Inches(6.5)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def create_report():
    doc = Document()
    
    # Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)

    # ================= COVER PAGE =================
    p_pre = doc.add_paragraph()
    p_pre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pre = p_pre.add_run("A Project Report\n")
    r_pre.font.size = Pt(14)
    r_pre.font.bold = True
    r_pre.font.color.rgb = RGBColor(100, 116, 139)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("CYBERLEARN")
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(31, 78, 121)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("A Scalable Hands-on Cybersecurity Learning & Practice Platform\n")
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(51, 65, 85)

    p_course = doc.add_paragraph()
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_course = p_course.add_run("CSE 400: Software Development Project IV\n")
    r_course.font.size = Pt(13)
    r_course.font.bold = True
    r_course.font.color.rgb = RGBColor(15, 23, 42)

    # SUBMITTED BY
    p_sb = doc.add_paragraph()
    r_sb = p_sb.add_run("SUBMITTED BY:")
    r_sb.font.size = Pt(12)
    r_sb.font.bold = True
    r_sb.font.color.rgb = RGBColor(31, 78, 121)

    tbl_team = doc.add_table(rows=5, cols=4)
    team_data = [
        ["Name", "ID", "Intake", "Section"],
        ["Muhammad Raihan", "22235103068", "51", "2"],
        ["Asif Raihan Rafi", "22235103281", "51", "2"],
        ["Ehsanul Haque", "22235103075", "51", "2"],
        ["Sadia Noor", "22235103286", "51", "2"]
    ]
    for r_idx, row in enumerate(tbl_team.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = team_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(10)
    style_table(tbl_team, [2.2, 1.8, 1.2, 1.2])

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SUPERVISED BY
    p_sup = doc.add_paragraph()
    r_sup_title = p_sup.add_run("SUPERVISED BY:\n")
    r_sup_title.font.bold = True
    r_sup_title.font.size = Pt(12)
    r_sup_title.font.color.rgb = RGBColor(31, 78, 121)
    
    r_sup_name = p_sup.add_run("Shampa Banik\n")
    r_sup_name.font.bold = True
    r_sup_name.font.size = Pt(11.5)
    r_sup_dept = p_sup.add_run("Lecturer, Department of Computer Science and Engineering\n")
    r_sup_dept.font.italic = True
    r_sup_dept.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # INSTITUTION
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dept = p_inst.add_run("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING\n")
    r_dept.font.bold = True
    r_dept.font.size = Pt(13)
    r_dept.font.color.rgb = RGBColor(31, 78, 121)
    
    r_uni = p_inst.add_run("BANGLADESH UNIVERSITY OF BUSINESS AND TECHNOLOGY (BUBT)\n")
    r_uni.font.bold = True
    r_uni.font.size = Pt(14)
    
    r_date = p_inst.add_run("27th August, 2026")
    r_date.font.size = Pt(11)

    doc.add_page_break()

    # ================= ABSTRACT =================
    h_abs = doc.add_heading("Abstract", level=1)
    h_abs.paragraph_format.space_before = Pt(10)
    h_abs.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph(
        "Cybersecurity education today is dominated by two extremes: theory-heavy classroom courses that "
        "leave learners unable to perform real security tasks, and steep, unguided Capture-The-Flag (CTF) "
        "platforms that assume prior command-line fluency and intimidate absolute beginners. CyberLearn "
        "was built to close this gap. It is a full-stack, browser-accessible learning and hands-on practice platform "
        "that interleaves structured video/reading lessons with disposable, isolated lab sandboxes, wraps the "
        "experience in a gamified progress system (XP, streaks, badges, and a global leaderboard), and embeds "
        "an AI Cyber Coach that a learner can consult for hints and error explanations without leaving the lesson."
    )
    doc.add_paragraph(
        "The platform is implemented as a decoupled two-tier web application. The frontend is a Next.js 16 "
        "(React 19) single-page application written in TypeScript, styled with Tailwind CSS v4 and managed with "
        "the Zustand state store. The backend is a modular FastAPI (Python 3.11+) REST service backed by a "
        "SQLAlchemy 2.0 ORM data layer that targets PostgreSQL 16 in production (SQLite as a zero-configuration "
        "development fallback). The system exposes twelve functional domains — authentication, user profiles & KYC "
        "verification, courses, interactive labs, AI coaching, community discussion, gamified leaderboards, "
        "certificates, exams, batch/cohort management, billing/subscriptions, and an administrative command center — "
        "through 95 URL paths and 119 HTTP endpoint operations secured with JWT-based authentication, role-based "
        "access control (Student, Instructor, Admin), and IP-based sliding-window rate limiting."
    )
    doc.add_paragraph(
        "This report documents the problem motivating the project, an analysis of existing platforms "
        "(TryHackMe, Hack The Box, and PortSwigger Web Security Academy), the chosen Agile/Scrum "
        "development model, the system's architecture and database design, the implementation of the front-end "
        "and back-end layers, a user manual describing every major screen of the deployed application, "
        "and a concluding discussion of the project's limitations and planned future work."
    )

    doc.add_page_break()

    # ================= LIST OF FIGURES =================
    h_lof = doc.add_heading("List of Figures", level=1)
    h_lof.paragraph_format.space_after = Pt(12)

    tbl_lof = doc.add_table(rows=17, cols=3)
    lof_data = [
        ["Sl. No.", "Figure Name", "Location"],
        ["1", "System Architecture Diagram", "Chapter 3"],
        ["2", "Use Case Diagram", "Chapter 3"],
        ["3", "Context Level Diagram (DFD-0)", "Chapter 3"],
        ["4", "Data Flow Diagram — Level 1", "Chapter 3"],
        ["5", "Entity–Relationship / Database Schema Diagram (20 Tables)", "Chapter 3"],
        ["6", "Lab Session Lifecycle & HMAC Verification Flowchart", "Chapter 3"],
        ["7", "Landing Page — Hero, Learning Paths & Interactive Tool Preview", "Chapter 5"],
        ["8", "Authentication Screens (Login, Signup, OAuth & Email Verification)", "Chapter 5"],
        ["9", "4-Step Student Onboarding Flow", "Chapter 5"],
        ["10", "Student Dashboard with Skill Radar Chart & Progress Metrics", "Chapter 5"],
        ["11", "Course Catalogue, Syllabus Detail & Interactive Lesson Player", "Chapter 5"],
        ["12", "Interactive Lab Workspace (Web Proxy, Terminal, SOC Logs, Socratic Hints)", "Chapter 5"],
        ["13", "Multi-Session AI Cyber Coach with Persona Switcher", "Chapter 5"],
        ["14", "Industry-Standard Certification Exams Catalogue", "Chapter 5"],
        ["15", "Global Leaderboard & Community Discussion Forum", "Chapter 5"],
        ["16", "Admin Command Center — User Governance & System Telemetry", "Chapter 5"]
    ]
    for r_idx, row in enumerate(tbl_lof.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = lof_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(10)
    style_table(tbl_lof, [0.8, 4.4, 1.3])

    doc.add_page_break()

    # ================= CHAPTER 1 =================
    doc.add_heading("Chapter 1: INTRODUCTION", level=1)
    
    doc.add_heading("1.1. Problem Specification", level=2)
    doc.add_paragraph(
        "Existing cybersecurity education resources exhibit critical limitations that hinder beginner onboarding "
        "and practical skill acquisition. Traditional classroom-style courses concentrate on lecture slides and "
        "multiple-choice quizzes, leaving students unable to actually operate a Linux terminal, exploit a "
        "vulnerable web form, or read a packet capture when they finish the course. On the opposite end, "
        "established Capture-The-Flag (CTF) platforms such as Hack The Box assume the learner already knows "
        "how to configure a VPN, use a hacking distribution, and navigate a shell — a level of comfort that most "
        "complete beginners do not have, which causes many of them to abandon the field early."
    )
    doc.add_paragraph(
        "In addition, building realistic lab environments traditionally requires local virtual machines, "
        "hypervisors, or paid cloud licences, creating a cost and setup barrier before a student has typed a "
        "single command. Finally, students who practise alone frequently get stuck on a syntax error or a logical "
        "misunderstanding with no instructor or personalised guide available to unblock them in real time."
    )
    doc.add_paragraph(
        "CyberLearn addresses these four problems together: it removes the local-setup barrier by provisioning "
        "disposable, isolated sandboxes directly in the browser; it removes the theory/practice gap by "
        "interleaving every lesson with an immediately-attached lab; it removes the cost barrier by using an "
        "open-source stack; and it removes the “stuck alone” problem with an integrated AI Cyber Coach that "
        "can explain a terminal error or nudge the learner toward the next step without giving away the full solution."
    )

    doc.add_heading("1.2. Objectives", level=2)
    doc.add_paragraph("The primary objectives of the CyberLearn project are to:")
    bullet_points = [
        "Provide disposable, fully isolated sandbox environments (Linux CLI, web security tools, SOC logs) that can be provisioned in seconds without any local software installation.",
        "Interleave structured reading and video lessons with immediate, hands-on lab challenges so that theory is reinforced by practice within the same session.",
        "Embed an AI Cyber Coach capable of generating contextual Socratic hints and explaining command or logic errors, reducing the need for a live human instructor.",
        "Increase learner retention through a gamified progression system — experience points (XP), daily streaks, unlockable badges, and a global leaderboard.",
        "Issue publicly verifiable digital certificates upon course and exam completion, giving learners a portfolio artefact they can share with employers.",
        "Incorporate a National ID (NID/KYC) identity verification pipeline to ensure the authenticity and industry credibility of issued credentials.",
        "Support role-based access (Student, Instructor, Admin) so that instructors can run cohort-based batches and administrators can manage content, users, and platform finances from a single control panel.",
        "Keep the entire system technically and financially feasible for a student-led team to design, build, and deploy within one academic term."
    ]
    for bp in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bp)

    doc.add_heading("1.3. Scope", level=2)
    p_in = doc.add_paragraph()
    r_in = p_in.add_run("In-Scope (Implemented & Shipped):")
    r_in.font.bold = True
    in_scope_items = [
        "Secure registration and login (email/password and Google/GitHub OAuth) with role-based access control for Student, Instructor, and Admin.",
        "20-table normalized relational schema backed by PostgreSQL/SQLite with automated multi-dialect migrations.",
        "Structured course catalogue with video, reading, and quiz lesson types, and per-lesson/per-course progress tracking.",
        "One-click lab session start/reset/submit flow with interactive Web Proxy Inspector, Network Topology Graph, and SOC Log Workbench.",
        "HMAC-SHA256 based dynamic flag generation and server-side verification.",
        "A floating multi-session AI Cyber Coach with persistent chat history and specialized persona prompts.",
        "XP, streaks, unlockable badges, and a global leaderboard.",
        "Timed certification exams with anti-cheat monitoring (fullscreen and blur detection) and automatically issued, publicly verifiable certificates.",
        "Government National ID (NID) verification upload and administrative review queue.",
        "A community discussion board (posts, comments, upvotes, solved/unsolved marking).",
        "Batch/cohort management for instructors (`CYBER-XXXXXX` join codes) and a comprehensive administrative command center.",
        "A billing module supporting course lifetime purchases and subscription plans (Free / Pro / Premium)."
    ]
    for item in in_scope_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    p_out = doc.add_paragraph()
    r_out = p_out.add_run("Out-of-Scope (Reserved for Future Work):")
    r_out.font.bold = True
    out_scope_items = [
        "Multi-user, real-time CTF Arenas with attack-defence scoring.",
        "Enterprise/organisational dashboards for corporate training teams.",
        "Locally hosted, offline LLM inference for the AI Coach (currently uses hosted Gemini/OpenAI API).",
        "A dedicated native mobile application."
    ]
    for item in out_scope_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_heading("1.4. Organization of Project Report", level=2)
    doc.add_paragraph(
        "The remainder of this report is organized as follows: Chapter 2 reviews background literature and existing systems; "
        "Chapter 3 presents the system analysis, Agile model, architecture, DFDs, ERD, and flag algorithm; Chapter 4 describes "
        "the frontend and backend implementation; Chapter 5 presents the user manual and live deployment; Chapter 6 concludes with "
        "limitations and future work, followed by References and Appendix A."
    )

    doc.add_page_break()

    # ================= CHAPTER 2 =================
    doc.add_heading("Chapter 2: BACKGROUND", level=1)
    
    doc.add_heading("2.1. Existing System Analysis", level=2)
    doc.add_paragraph(
        "Three widely used cybersecurity learning platforms were studied to identify functional gaps that CyberLearn could address:"
    )
    doc.add_paragraph(
        "A. TryHackMe: TryHackMe is a gamified, room-based platform that provides web-based target boxes through a split-screen "
        "browser interface. It is beginner-friendly and uses XP/badges to drive engagement, which directly inspired CyberLearn's "
        "own gamification system. However, its subscription pricing is relatively costly for students in developing economies, and "
        "the platform lacks an adaptive AI assistant that can help a stuck user debug a specific command or syntax error in real time."
    )
    doc.add_paragraph(
        "B. Hack The Box (HTB): HTB is a competitive penetration-testing platform offering raw server access for exploitation tasks. "
        "Its content is technically excellent, but the learning curve is very steep for absolute beginners: there is minimal guided-path "
        "assistance, and users are expected to configure their own local VPN connections and attack tooling before they can even reach a target."
    )
    doc.add_paragraph(
        "C. PortSwigger Web Security Academy: PortSwigger's Web Security Academy is an excellent free resource for learning HTTP-based "
        "vulnerabilities in depth. Its scope, however, is limited exclusively to web application security; it offers no general Linux "
        "command-line practice, no packet-sniffing/networking labs, and no OSINT or reconnaissance exercises."
    )
    doc.add_paragraph(
        "Gap Identified: No single low-cost platform combines a genuinely beginner-friendly onboarding experience, multi-domain labs "
        "(Linux, networking, web, cloud, blue-team, and AI security), cohort-based batch management, and an integrated Socratic AI tutor. "
        "CyberLearn was designed specifically to occupy that gap."
    )

    doc.add_heading("2.2. Supporting Literatures & Methodologies", level=2)
    lit_points = [
        "Client–Server / REST Architecture: The platform follows the classical two-tier client–server model, with a stateless REST API separating presentation (Next.js) from business logic and persistence (FastAPI + PostgreSQL).",
        "Relational Database Theory: The schema was normalised using standard 3NF entity–relationship modelling across 20 tables so that user, course, lab, and financial data maintain referential integrity through foreign keys and unique constraints.",
        "Stateless Authentication: JSON Web Tokens (JWT, RFC 7519) signed with HS256 allow the FastAPI backend to verify user identity on every request without a server-side session store.",
        "Cryptographic Flag Verification: Lab flags are derived using HMAC-SHA256 (a keyed-hash message authentication code) rather than being stored in plaintext in the database.",
        "Agile/Iterative Delivery: An iterative, sprint-based delivery model was adopted to mitigate integration risks across decoupled frontend and backend layers."
    ]
    for lp in lit_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(lp)

    doc.add_page_break()

    # ================= CHAPTER 3 =================
    doc.add_heading("Chapter 3: SYSTEM ANALYSIS & DESIGN", level=1)
    
    doc.add_heading("3.1. Technology & Tools", level=2)
    doc.add_paragraph(
        "CyberLearn is built as a decoupled two-tier web application: a Next.js single-page frontend communicating over a "
        "versioned REST API with a modular FastAPI backend. Table 1 summarises the complete technology stack used in the shipped codebase."
    )

    tbl_tech = doc.add_table(rows=18, cols=2)
    tech_data = [
        ["Layer / Concern", "Technology Used"],
        ["Frontend Framework", "Next.js 16.2.9 (App Router) with React 19, TypeScript 5"],
        ["Styling & Theme", "Tailwind CSS v4 (Dark-first NetAcad Design System)"],
        ["Frontend State Management", "Zustand 5"],
        ["Forms & Validation", "React Hook Form 7 + Zod 4 resolvers"],
        ["Animation & Icons", "Framer Motion 12, Lucide React"],
        ["Interactive Lab Tools", "WebProxyInspector, NetworkTopologyGraph, SocLogWorkbench, SocraticHintDrawer"],
        ["Backend Framework", "FastAPI (Python 3.11+), Uvicorn ASGI server"],
        ["ORM / Data Layer", "SQLAlchemy 2.0"],
        ["Database (Production)", "PostgreSQL 16"],
        ["Database (Development)", "SQLite (zero-config local fallback with auto-migrations)"],
        ["Authentication", "JWT (PyJWT / HS256) + Passlib/bcrypt + Firebase Auth SDK"],
        ["Rate Limiting", "In-Memory Sliding-Window ASGI Middleware"],
        ["AI Coach LLM Provider", "Google Gemini API (Gemini 3.5 Flash / Pro) & OpenAI API"],
        ["OAuth Providers", "Google OAuth 2.0, GitHub OAuth"],
        ["Email Delivery", "SMTP / Brevo (Sendinblue) transactional API"],
        ["Hosting / Deployment", "Vercel (frontend edge) + Render/Docker host (backend API)"],
        ["Version Control", "Git & GitHub"]
    ]
    for r_idx, row in enumerate(tbl_tech.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = tech_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9.5)
    style_table(tbl_tech, [2.5, 4.0])

    p_cap1 = doc.add_paragraph()
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c1 = p_cap1.add_run("Table 1: Technology stack used in CyberLearn")
    r_c1.font.italic = True
    r_c1.font.size = Pt(9.5)

    doc.add_heading("3.2. Model & Diagram", level=2)
    doc.add_heading("3.2.1. Model (SDLC — Agile/Scrum)", level=3)
    doc.add_paragraph(
        "The team selected the Iterative and Incremental (Agile/Scrum) model to guide development across four short sprints:\n"
        "• Sprint 1 — Core data models, authentication (register/login/JWT/OAuth), and the course catalogue API.\n"
        "• Sprint 2 — Interactive lab tools (Web Proxy, SOC Logs, Network Graph), and HMAC-SHA256 flag verification.\n"
        "• Sprint 3 — Gamification (XP/streaks/badges/leaderboard), multi-session AI Cyber Coach, and certification exams.\n"
        "• Sprint 4 — Community discussion, instructor batches, KYC National ID review queue, and administrative command center."
    )

    doc.add_heading("3.2.2. Database Schema (20 Tables)", level=3)
    doc.add_paragraph(
        "CyberLearn uses a normalised relational schema of 20 tables managed through SQLAlchemy 2.0. Every table uses a UUID "
        "primary key, and foreign keys enforce referential integrity with cascading deletes. Table 2 outlines the core entity groups."
    )

    tbl_db = doc.add_table(rows=12, cols=2)
    db_data = [
        ["Entity / Group", "Purpose in System"],
        ["users", "Account credentials, roles (Student, Instructor, Admin), XP/streak, KYC verification status, and subscription tier."],
        ["courses / lessons", "Structured cybersecurity catalogue: course metadata, pricing, video URLs, markdown notes, and quizzes."],
        ["progress", "Per-user, per-lesson completion tracking (unique constraint on user_id + course_id + lesson_id)."],
        ["labs / lab_sessions", "Practice lab specifications, container templates, and live student session timers/states."],
        ["achievements", "Timestamped milestone badges earned per student."],
        ["ai_sessions / ai_chat_messages", "Persistent, multi-session AI Coach conversation history with custom persona system prompts."],
        ["exams / exam_questions / exam_submissions", "Timed assessments, 20-25 MCQ banks, anti-cheat tracking, and graded student submissions."],
        ["certificates", "Publicly verifiable completion and certification credentials linked by unique SHA-256 tokens."],
        ["posts / comments / post_votes", "Community discussion threads, threaded replies, solution markings, and upvote ledger."],
        ["batches / batch_enrollments", "Instructor-led cohorts (`CYBER-XXXXXX`), schedules, rosters, and seat capacities."],
        ["invoices / course_purchases", "Financial audit ledger for subscription plans and lifetime individual course purchases."]
    ]
    for r_idx, row in enumerate(tbl_db.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = db_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9.5)
    style_table(tbl_db, [2.2, 4.3])

    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c2 = p_cap2.add_run("Table 2: Core database entities and their purpose")
    r_c2.font.italic = True
    r_c2.font.size = Pt(9.5)

    doc.add_heading("3.2.3. Algorithm / Flowchart", level=3)
    doc.add_paragraph(
        "The platform's flag-verification workflow uses deterministic HMAC-SHA256 hashing. The exact algorithm from "
        "backend/app/labs/router.py is provided below:"
    )
    add_code_block(doc, 
        "def generate_lab_flag(lab_id: str) -> str:\n"
        "    signature = hmac.new(\n"
        "        settings.FLAG_SECRET.encode('utf-8'),\n"
        "        lab_id.encode('utf-8'),\n"
        "        hashlib.sha256\n"
        "    ).hexdigest()[:16]\n"
        "    return f'FLAG{{{signature}}}'\n"
        "\n"
        "# Flag is verified server-side on submission without database plaintext storage"
    )

    doc.add_page_break()

    # ================= CHAPTER 4 =================
    doc.add_heading("Chapter 4: IMPLEMENTATION", level=1)
    
    doc.add_heading("4.1. Interface Design / Front-End", level=2)
    doc.add_paragraph(
        "The frontend is a Next.js 16 application using the App Router with TypeScript and Tailwind CSS v4. Table 3 lists "
        "the primary routes implemented."
    )

    tbl_routes = doc.add_table(rows=16, cols=2)
    routes_data = [
        ["Route", "Purpose"],
        ["/ (Landing)", "Public marketing landing page with hero, interactive tools preview, and pricing."],
        ["/(auth)/login, /(auth)/signup", "Authentication screens with email/password and Google/GitHub OAuth."],
        ["/(app)/dashboard", "Personalised student dashboard — stats, continue-learning, and skill radar chart."],
        ["/(app)/courses, /courses/[courseId]", "Course catalogue, module outline, and interactive video/quiz lesson player."],
        ["/(app)/labs, /labs/[labId]", "Lab workspace with Web Proxy Inspector, SOC Logs, and Socratic hints."],
        ["/(app)/challenges", "Standalone challenge/CTF-style exercise browser."],
        ["/(app)/ai-coach", "Full-page AI Cyber Coach chat interface with multi-session history."],
        ["/(app)/exams, /exams/[examId]", "Timed certification exam listing and anti-cheat monitored attempt flow."],
        ["/(app)/verify-nid", "Government National ID (NID/Passport) identity verification upload portal."],
        ["/(app)/certificates", "Student's earned certificates with download, share, and QR actions."],
        ["/(app)/leaderboard", "Global XP leaderboard with podium rankings."],
        ["/(app)/community, /community/[postId]", "Community discussion feed, post detail, and threaded comments."],
        ["/(app)/batches, /batches/[batchCode]", "Instructor cohort management and batch roster/schedule view."],
        ["/(app)/admin", "Full administrative command center (users, KYC, courses, labs, financials)."],
        ["/portfolio/[username]", "Public shareable learner portfolio and credential showcase."],
        ["/verify/[token]", "Standalone public certificate authenticity verification portal."]
    ]
    for r_idx, row in enumerate(tbl_routes.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = routes_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9.5)
    style_table(tbl_routes, [2.5, 4.0])

    p_cap3 = doc.add_paragraph()
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c3 = p_cap3.add_run("Table 3: Primary frontend routes implemented")
    r_c3.font.italic = True
    r_c3.font.size = Pt(9.5)

    doc.add_heading("4.2. Back-End", level=2)
    doc.add_paragraph(
        "The backend is a FastAPI application composed of 12 independent, prefix-routed modules registered onto a single "
        "FastAPI app instance in main.py. Table 4 summarises the responsibilities of each router module."
    )

    tbl_backend = doc.add_table(rows=13, cols=2)
    be_data = [
        ["Router Module", "Responsibility"],
        ["auth", "Registration, email verification, login, logout, OAuth code exchange, and username checks."],
        ["users", "Profile read/update, avatar upload, password changes, public profiles, and NID KYC submission."],
        ["courses", "Course listing, syllabus delivery, lesson progress, and quiz grading."],
        ["labs", "Lab catalogue, session start/reset/submit, HMAC flag verification, hint unlocking, and Web Proxy forwarding."],
        ["ai", "AI Coach chat (stateless & session-based), session CRUD, and conversation history."],
        ["exams", "Exam catalog, question bank delivery, timed submission grading, and submission history."],
        ["certificates", "Student certificate listing and public token-based authenticity verification."],
        ["leaderboard", "Global XP ranking and user rank positioning."],
        ["community", "Post CRUD, threaded comments, upvoting, and solved/unsolved toggling."],
        ["batches", "Batch creation, student joining by code (`CYBER-XXXXXX`), and roster tracking."],
        ["billing", "Promo validation, payment checkout, invoice history, and subscription lifecycle."],
        ["admin", "Full CRUD governance over users, KYC queues, courses, exams, labs, and revenue telemetry."]
    ]
    for r_idx, row in enumerate(tbl_backend.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = be_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9.5)
    style_table(tbl_backend, [1.8, 4.7])

    p_cap4 = doc.add_paragraph()
    p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c4 = p_cap4.add_run("Table 4: Backend router modules and their responsibilities")
    r_c4.font.italic = True
    r_c4.font.size = Pt(9.5)

    doc.add_page_break()

    # ================= CHAPTER 5 =================
    doc.add_heading("Chapter 5: USER MANUAL", level=1)
    
    doc.add_heading("5.1. System Requirement", level=2)
    doc.add_paragraph(
        "Client (End User):\n"
        "• Processor: Dual-core 1.6 GHz+ processor\n"
        "• RAM: 4 GB minimum (8 GB recommended for smooth multi-tab use)\n"
        "• Operating System: Windows 10/11, macOS 12+, Linux, Android, or iOS\n"
        "• Browser: Google Chrome, Microsoft Edge, or Mozilla Firefox\n\n"
        "Server (Deployment Host):\n"
        "• Host OS: Ubuntu 22.04 LTS\n"
        "• Runtime: Python 3.11+, Node.js 18+, Docker Engine 24+ & Docker Compose v2\n"
        "• Database: PostgreSQL 16 (Production) / SQLite (Local development)"
    )

    doc.add_heading("5.2. Live Deployment & Credentials", level=2)
    doc.add_paragraph(
        "CyberLearn is deployed and publicly accessible. Seeded demonstration accounts are provided below:"
    )

    tbl_deploy = doc.add_table(rows=5, cols=2)
    dep_data = [
        ["Item", "Value"],
        ["Deployed Application URL", "https://cyber-learn-three.vercel.app"],
        ["Source Repository", "https://github.com/raihan12121/CyberLearn"],
        ["Demo Admin Login", "admin@cyberlearn.io / admin123"],
        ["Demo Student Login", "learner@cyberlearn.io / learner123"]
    ]
    for r_idx, row in enumerate(tbl_deploy.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = dep_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            if r_idx == 0:
                p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(10)
    style_table(tbl_deploy, [2.5, 4.0])

    p_cap5 = doc.add_paragraph()
    p_cap5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c5 = p_cap5.add_run("Table 5: Deployment link and demo login credentials")
    r_c5.font.italic = True
    r_c5.font.size = Pt(9.5)

    doc.add_page_break()

    # ================= CHAPTER 6 =================
    doc.add_heading("Chapter 6: CONCLUSION & FUTURE WORK", level=1)
    
    doc.add_heading("6.1. Conclusion", level=2)
    doc.add_paragraph(
        "CyberLearn successfully delivers on its core objective: a single, low-cost, browser-accessible platform that takes "
        "a complete beginner from structured lessons directly into hands-on practice, backed by an AI Cyber Coach, a 20-table "
        "normalized relational schema, and 119 API endpoint operations across 12 domain routers. The system provides an end-to-end "
        "learning lifecycle from onboarding to recruiter-grade KYC verified certification."
    )

    doc.add_heading("6.2. Limitations", level=2)
    doc.add_paragraph(
        "• External LLM Metering: Relies on external hosted APIs (Google Gemini / OpenAI); heavy traffic incurs metered per-token API costs.\n"
        "• Sandbox Egress Lock: Practice sandboxes are denied outbound networking to prevent malicious abuse, which prevents students from installing arbitrary external tools inside running labs.\n"
        "• Concurrent Capacity: Simultaneous container sessions are bounded by server memory/CPU limits."
    )

    doc.add_heading("6.3. Planned Future Works", level=2)
    doc.add_paragraph(
        "• Multiplayer Real-Time CTF Arenas with attack-defence live scoring.\n"
        "• On-premise local open-weight LLM inference (e.g. DeepSeek / Llama 3) to eliminate external API costs.\n"
        "• Kubernetes pod orchestration for elastic multi-node scaling.\n"
        "• Native iOS and Android mobile applications."
    )

    doc.add_page_break()

    # ================= REFERENCES =================
    doc.add_heading("References", level=1)
    refs = [
        "[1] TryHackMe Ltd. 'Interactive Cybersecurity Training.' TryHackMe, 2024, tryhackme.com.",
        "[2] Hack The Box. 'Gamified Cybersecurity Practice Labs.' Hack The Box, 2024, www.hackthebox.com.",
        "[3] PortSwigger. 'Web Security Academy.' PortSwigger Ltd., 2024, portswigger.net/web-security.",
        "[4] OWASP Foundation. 'OWASP Top 10 Web Application Security Risks.' OWASP, 2021, owasp.org/www-project-top-ten.",
        "[5] Docker Inc. 'Docker Container Isolation and Security Best Practices.' Docker Documentation, 2023, docs.docker.com/engine/security.",
        "[6] Internet Engineering Task Force. 'RFC 7519: JSON Web Token (JWT).' IETF, May 2015, datatracker.ietf.org/doc/html/rfc7519.",
        "[7] Vercel Inc. 'Next.js Documentation.' Next.js, 2026, nextjs.org/docs.",
        "[8] FastAPI. 'FastAPI Documentation.' FastAPI, 2026, fastapi.tiangolo.com.",
        "[9] SQLAlchemy. 'SQLAlchemy 2.0 Documentation.' SQLAlchemy, 2026, docs.sqlalchemy.org.",
        "[10] PostgreSQL Global Development Group. 'PostgreSQL 16 Documentation.' PostgreSQL, 2026, www.postgresql.org/docs/16."
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(ref)
        r.font.size = Pt(10)

    doc.add_page_break()

    # ================= APPENDIX A =================
    doc.add_heading("Appendix A", level=1)
    
    doc.add_heading("A.1. Project Repository & Deployment", level=2)
    doc.add_paragraph(
        "Source code repository: https://github.com/raihan12121/CyberLearn\n"
        "Live deployment: https://cyber-learn-three.vercel.app\n"
        "Backend interactive API documentation (Swagger/OpenAPI): Available at backend host `/docs` path."
    )

    doc.add_heading("A.2. Selected Directory Structure", level=2)
    add_code_block(doc,
        "CyberLearn/\n"
        "├── backend/                 # FastAPI backend\n"
        "│   ├── app/\n"
        "│   │   ├── auth/ users/ courses/ labs/ ai/\n"
        "│   │   ├── community/ leaderboard/ certificates/\n"
        "│   │   ├── billing/ batches/ exams/ admin/\n"
        "│   │   ├── database.py models.py schemas.py main.py\n"
        "│   ├── requirements.txt\n"
        "│   └── test_*.py            # Comprehensive Pytest suites\n"
        "├── frontend/                # Next.js 16 application\n"
        "│   ├── src/app/(auth)       # /login, /signup\n"
        "│   ├── src/app/(app)/...    # /dashboard, /courses, /labs, /admin\n"
        "│   ├── src/app/verify/      # Standalone public certificate verifier\n"
        "│   ├── src/components/      # WebProxy, SocLogs, TopologyGraph, UI library\n"
        "│   ├── src/lib/             # authStore.ts, api.ts\n"
        "│   └── package.json\n"
        "├── docker-compose.yml\n"
        "├── README.md\n"
        "└── PRD.md TRD.md DesignDescription.md ProjectProposal.md"
    )

    doc.add_heading("A.3. Automated Test Coverage", level=2)
    doc.add_paragraph(
        "The repository includes comprehensive automated test suites exercising all critical workflows:\n"
        "• Root Test Suites: test_admin_control_panel.py, test_id_verification_and_20_mcqs.py, test_tasks_1_2_3.py, "
        "test_exam_and_course_completion_storage.py, test_google_avatar_and_profile_updates.py, test_login_session_persistence.py, "
        "test_onboarding_and_unique_username.py.\n"
        "• Backend Pytest Suites: test_api_endpoints.py, test_blackbox_workflows.py, test_qa_regressions.py, "
        "test_subscription_access.py, test_whitebox_units.py, test_hybrid_revenue.py, test_realistic_checkout.py."
    )

    output_path = r"d:\web app\CyberLearn\CyberLearn_Project_Report.docx"
    doc.save(output_path)
    print(f"[SUCCESS] Document generated successfully at: {output_path}")

if __name__ == "__main__":
    create_report()
